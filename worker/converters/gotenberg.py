import os
import time
import subprocess
import markdown
import httpx
from pathlib import Path
from worker.config import get_settings

settings = get_settings()

GOTENBERG_URL = settings.GOTENBERG_URL
LIBREOFFICE_ROUTE = "/forms/libreoffice/convert"
PDFENGINES_ROUTE  = "/forms/pdfengines/convert"

# Gotenberg two-pass temp directory — must match the docker-compose volume mount.
# The worker and gotenberg containers share this volume so Gotenberg can read
# the intermediate PDF produced in step 1 during step 2.
GOTENBERG_SHARED_DIR = "/tmp/gotenberg"

MIME_MAP = {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "html": "text/html",
    "txt":  "text/plain",
    "png":  "image/png",
    "jpg":  "image/jpeg",
    "jpeg": "image/jpeg",
    "odt":  "application/vnd.oasis.opendocument.text",
    "ods":  "application/vnd.oasis.opendocument.spreadsheet",
    "odp":  "application/vnd.oasis.opendocument.presentation",
    "epub": "application/epub+zip",
    "rtf":  "application/rtf",
    "doc":  "application/msword",
    "xls":  "application/vnd.ms-excel",
    "ppt":  "application/vnd.ms-powerpoint",
    "csv":  "text/csv",
}

# Image formats that Gotenberg LibreOffice does NOT support directly.
# These must be pre-converted to a supported format before sending to Gotenberg.
_UNSUPPORTED_BY_GOTENBERG = frozenset(("avif", "webp", "tiff", "gif"))

# Gotenberg-supported image formats (for pre-conversion before avif→pdf etc.)
_GOTENBERG_IMAGE_FORMATS = ("png", "jpg", "jpeg", "bmp")


def convert(
    input_path: str,
    output_path: str,
    input_format: str,
    output_format: str,
    timeout_s: float = 120.0,
) -> None:
    from worker.security.path_guard import sanitize_and_assert_tmp_path
    from PIL import Image

    sanitize_and_assert_tmp_path(input_path)
    sanitize_and_assert_tmp_path(output_path)

    # ── Pre-process: convert unsupported image formats to Gotenberg-supported ──
    # AVIF, WebP, TIFF, GIF cannot be sent directly to Gotenberg LibreOffice.
    # Pre-convert to BMP (which Gotenberg supports) using Pillow first.
    preconverted_path = input_path
    if input_format in _UNSUPPORTED_BY_GOTENBERG:
        tmp_dir = Path("/tmp/got-preconv")
        tmp_dir.mkdir(exist_ok=True)
        # Convert to BMP — Gotenberg's LibreOffice accepts BMP natively
        bmp_path = tmp_dir / f"preconv.{os.path.basename(input_path)}.bmp"
        img = Image.open(input_path)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.save(str(bmp_path), format="BMP")
        preconverted_path = str(bmp_path)
        print(f"[gotenberg] pre-converted {input_format} -> BMP for Gotenberg")

    # Route selection:
    #   - PDF → PNG/JPG: use pdf2image (more reliable than Gotenberg)
    #   - All others: LIBREOFFICE_ROUTE (handles PNG/JPG→PDF too)
    endpoint = LIBREOFFICE_ROUTE

    if input_format == "pdf" and output_format in ("png", "jpg", "jpeg"):
        # BUG 1 fix: PDF → image via pdf2image instead of Gotenberg
        _pdf_to_image(input_path, output_path, output_format, timeout_s)
        return

    # Build file for upload
    file_bytes: bytes
    mime: str
    ext: str
    if input_format == "md":
        with open(preconverted_path, "r", encoding="utf-8") as f:
            md_text = f.read()
        html_content = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
        file_bytes = html_content.encode("utf-8")
        mime = "text/html"
        ext = "html"
    else:
        with open(preconverted_path, "rb") as f:
            file_bytes = f.read()
        mime = MIME_MAP.get(input_format, "application/octet-stream")
        ext = Path(preconverted_path).suffix.lstrip(".") or input_format

    start = time.perf_counter()

    with httpx.Client(timeout=httpx.Timeout(timeout_s, connect=5.0)) as client:
        response = client.post(
            f"{GOTENBERG_URL}{endpoint}",
            files={
                "files": (f"input.{ext}", file_bytes, mime),
                # BUG 1 fix: Gotenberg LibreOffice requires nativePageRanges
                "nativePageRanges": (None, "1-"),
            },
        )

    duration_ms = int((time.perf_counter() - start) * 1000)
    print(f"[gotenberg] converting {input_format}->{output_format} via {endpoint} in {duration_ms}ms")

    if response.status_code != 200:
        raise RuntimeError(
            f"gotenberg: HTTP {response.status_code} for "
            f"{input_format}->{output_format}: {response.text[:500]}"
        )

    if output_format == "pdf":
        _write_output(response, output_path)
        return

    # ── Two-pass conversion (BUG 1 fix) ────────────────────────────────────
    #
    # Gotenberg's LibreOffice endpoint always returns PDF. For non-PDF outputs:
    #   step 1: save Gotenberg's PDF response to shared volume
    #   step 2: post-process the PDF to the target format
    #
    # step 2 helpers:
    #   txt  → pdftotext (most reliable for plain text)
    #   docx → pdftotext → python-docx document
    #   csv  → pdftotext → best-effort CSV parsing
    #   html → pdftotext -htmlmeta OR minimal HTML wrapper
    #
    os.makedirs(GOTENBERG_SHARED_DIR, exist_ok=True)
    tmp_pdf = os.path.join(GOTENBERG_SHARED_DIR, f"got-tmp-{os.getpid()}-{time.time_ns()}.pdf")

    try:
        with open(tmp_pdf, "wb") as f:
            for chunk in response.iter_bytes(chunk_size=8192):
                f.write(chunk)

        if output_format == "txt":
            _pdf_to_txt(tmp_pdf, output_path, timeout_s)
        elif output_format == "docx":
            _pdf_to_docx(tmp_pdf, output_path, timeout_s)
        elif output_format == "csv":
            _pdf_to_csv(tmp_pdf, output_path, timeout_s)
        elif output_format == "html":
            _pdf_to_html(tmp_pdf, output_path, timeout_s)
        else:
            # Fallback: treat as docx
            _pdf_to_docx(tmp_pdf, output_path, timeout_s)
    finally:
        if os.path.exists(tmp_pdf):
            os.unlink(tmp_pdf)

    if os.path.getsize(output_path) == 0:
        raise RuntimeError(f"gotenberg: empty output for {input_format}->{output_format}")

    print(f"[gotenberg] done in {duration_ms}ms ({os.path.getsize(output_path)} bytes)")


def _write_output(response: httpx.Response, output_path: str) -> None:
    try:
        with open(output_path, "wb") as f:
            for chunk in response.iter_bytes(chunk_size=8192):
                f.write(chunk)
    except Exception:
        if os.path.exists(output_path):
            os.unlink(output_path)
        raise


# ── PDF → target format helpers ────────────────────────────────────────────

def _pdf_to_image(
    pdf_path: str,
    output_path: str,
    output_format: str,
    timeout_s: float,
) -> None:
    """
    BUG 1 fix: use pdf2image (pdftoppm + PIL) to render PDF pages as raster
    images. Much more reliable than Gotenberg pdfengines for PNG/JPG output.
    """
    from pdf2image import convert_from_path
    from PIL import Image

    images = convert_from_path(
        pdf_path,
        dpi=150,
        first_page=1,
        last_page=1,
        fmt=output_format.upper(),
    )
    if not images:
        raise RuntimeError(f"[gotenberg] pdf2image produced no pages for {pdf_path}")
    img = images[0]
    img.save(output_path, format=FMT_MAP[output_format])
    print(f"[gotenberg] rendered PDF -> {output_format} via pdf2image")


# Map from output format string to PIL format name
FMT_MAP = {
    "png":  "PNG",
    "jpg":  "JPEG",
    "jpeg": "JPEG",
    "webp": "WEBP",
    "avif": "AVIF",
    "bmp":  "BMP",
    "tiff": "TIFF",
    "gif":  "GIF",
}


def _pdf_to_txt(pdf_path: str, output_path: str, timeout_s: float) -> None:
    """Extract plain text from PDF using pdftotext."""
    result = subprocess.run(
        ["pdftotext", "-layout", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
        check=True,
    )
    Path(output_path).write_bytes(result.stdout)


def _pdf_to_docx(pdf_path: str, output_path: str, timeout_s: float) -> None:
    """
    BUG 1 fix: build a minimal DOCX from the pdftotext output using python-docx.
    """
    result = subprocess.run(
        ["pdftotext", "-layout", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
        check=True,
    )
    text = result.stdout  # already bytes from capture_output

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for para_text in text.split(b"\n\n"):
        if isinstance(para_text, bytes):
            para_text = para_text.decode("utf-8", errors="replace")
        para_text = para_text.strip()
        if not para_text:
            continue
        para = doc.add_paragraph()
        run = para.add_run(para_text)
        run.font.size = Pt(11)

    doc.save(output_path)


def _pdf_to_csv(pdf_path: str, output_path: str, timeout_s: float) -> None:
    """
    Best-effort CSV from PDF: extract text via pdftotext and parse as CSV.
    Real CSV conversion requires the original XLSX source.
    """
    result = subprocess.run(
        ["pdftotext", "-layout", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
        check=True,
    )
    text = result.stdout
    if isinstance(text, bytes):
        text = text.decode("utf-8", errors="replace")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    csv_rows = [",".join(f'"{f}"' for f in l.split()) for l in lines]
    Path(output_path).write_text("\n".join(csv_rows), encoding="utf-8")


def _pdf_to_html(pdf_path: str, output_path: str, timeout_s: float) -> None:
    """
    Try pdftotext -htmlmeta first; if that fails, build a minimal HTML wrapper.
    """
    result = subprocess.run(
        ["pdftotext", "-htmlmeta", "-q", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
    )
    if result.returncode == 0 and result.stdout.strip():
        Path(output_path).write_bytes(result.stdout)
    else:
        text_result = subprocess.run(
            ["pdftotext", "-layout", pdf_path, "-"],
            capture_output=True,
            timeout=timeout_s,
        )
        text = text_result.stdout
        if isinstance(text, bytes):
            text = text.decode("utf-8", errors="replace")
        lines = ["<p>{}</p>".format(p.strip()) for p in text.split("\n\n") if p.strip()]
        html = "<!DOCTYPE html>\n<html><head><meta charset=\"utf-8\"/>\n<body>\n"
        html += "\n".join(lines)
        html += "\n</body></html>"
        Path(output_path).write_text(html, encoding="utf-8")
