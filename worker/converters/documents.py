# worker/converters/documents.py
# All document conversions.
# Libraries: Gotenberg (primary), Pandoc (md/rtf), pdf2docx (pdf→docx)
#
# Internal routing:
#   pdf→docx          → pdf2docx (best layout preservation)
#   md/rtf input     → Pandoc
#   docx→rtf          → Pandoc
#   everything else   → Gotenberg (LibreOffice)

import os
import time
import subprocess
import markdown
import httpx
import logging
from pathlib import Path

from worker.security.path_guard import sanitize_and_assert_tmp_path
from worker.config import get_settings

logger = logging.getLogger(__name__)

settings = get_settings()
GOTENBERG_URL = settings.GOTENBERG_URL

# Gotenberg two-pass shared directory — must match docker-compose volume mount.
# Worker and Gotenberg containers share this so Gotenberg can read the
# intermediate PDF produced in step 1 during step 2.
GOTENBERG_SHARED_DIR = "/tmp/gotenberg"

# Gotenberg LibreOffice returns these formats directly (via "format" param).
_GOTENBERG_DIRECT_FORMATS = frozenset(("docx", "xlsx", "pptx", "odt", "ods", "odp"))

# Formats Pandoc handles better than Gotenberg.
PANDOC_FORMATS = frozenset(("md", "rtf"))

MIME_MAP = {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "html": "text/html",
    "txt":  "text/plain",
    "odt":  "application/vnd.oasis.opendocument.text",
    "ods":  "application/vnd.oasis.opendocument.spreadsheet",
    "odp":  "application/vnd.oasis.opendocument.presentation",
    "rtf":  "application/rtf",
    "csv":  "text/csv",
}

# Image formats Gotenberg does NOT support directly — pre-convert to BMP.
_UNSUPPORTED_BY_GOTENBERG = frozenset(("avif", "webp", "tiff", "gif"))


def convert(
    input_path: str,
    output_path: str,
    input_format: str,
    output_format: str,
    timeout_s: float = 120.0,
) -> None:
    sanitize_and_assert_tmp_path(input_path)
    sanitize_and_assert_tmp_path(output_path)

    # pdf→docx: use pdf2docx for best layout preservation
    if input_format == "pdf" and output_format == "docx":
        logger.info(f"[{input_format}→{output_format}] using pdf2docx")
        _pdf_to_docx(input_path, output_path)

    # md/rtf input: use Pandoc
    elif input_format in PANDOC_FORMATS:
        logger.info(f"[{input_format}→{output_format}] using pandoc")
        _pandoc(input_path, output_path, input_format, output_format, timeout_s)

    # docx→rtf: Pandoc handles better than Gotenberg
    elif input_format == "docx" and output_format == "rtf":
        logger.info(f"[{input_format}→{output_format}] using pandoc")
        _pandoc(input_path, output_path, input_format, output_format, timeout_s)

    # everything else: Gotenberg
    else:
        logger.info(f"[{input_format}→{output_format}] using gotenberg")
        _gotenberg(input_path, output_path, input_format, output_format, timeout_s)

    _assert_output(output_path, input_format, output_format)


def _assert_output(output_path, input_format, output_format) -> None:
    p = Path(output_path)
    if not p.exists() or p.stat().st_size == 0:
        raise RuntimeError(
            f"documents converter produced no output for {input_format}→{output_format}"
        )


def _pdf_to_docx(input_path: str, output_path: str) -> None:
    from pdf2docx import Converter as PDF2DOCXConverter

    cv = PDF2DOCXConverter(input_path)
    try:
        cv.convert(output_path, start=0, end=None)
    finally:
        cv.close()


def _pandoc(
    input_path: str,
    output_path: str,
    input_format: str,
    output_format: str,
    timeout_s: float,
) -> None:
    result = subprocess.run(
        ["pandoc", input_path, "-o", output_path, "--standalone"],
        capture_output=True,
        text=True,
        timeout=timeout_s,
    )
    if result.returncode != 0:
        raise RuntimeError(f"Pandoc failed: {result.stderr[-300:]}")


# ── Gotenberg internals (copied from existing gotenberg.py) ─────────────────

from PIL import Image


def _gotenberg(
    input_path: str,
    output_path: str,
    input_format: str,
    output_format: str,
    timeout_s: float,
) -> None:
    # Pre-process: convert unsupported image formats to Gotenberg-supported BMP.
    preconverted_path = input_path
    if input_format in _UNSUPPORTED_BY_GOTENBERG:
        tmp_dir = Path("/tmp/got-preconv")
        tmp_dir.mkdir(exist_ok=True)
        bmp_path = tmp_dir / f"preconv.{os.path.basename(input_path)}.bmp"
        img = Image.open(input_path)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.save(str(bmp_path), format="BMP")
        preconverted_path = str(bmp_path)
        logger.info(f"[gotenberg] pre-converted {input_format} -> BMP for Gotenberg")

    # PDF → PNG/JPG: use pdf2image (more reliable than Gotenberg)
    if input_format == "pdf" and output_format in ("png", "jpg", "jpeg"):
        _gotenberg_pdf_to_image(input_path, output_path, output_format, timeout_s)
        return

    # Build file for upload
    if input_format == "md":
        with open(preconverted_path, "r", encoding="utf-8") as f:
            md_text = f.read()
        html_content = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
        file_bytes = html_content.encode("utf-8")
        mime = "text/html"
        ext = "html"
    else:
        with open(preconverted_path, "rb") as f:
            file_bytes = f.read()
        mime = MIME_MAP.get(input_format, "application/octet-stream")
        ext = Path(preconverted_path).suffix.lstrip(".") or input_format

    start = time.perf_counter()

    form_fields = [("files", (f"input.{ext}", file_bytes, mime))]
    if output_format in _GOTENBERG_DIRECT_FORMATS:
        form_fields.append(("format", (None, output_format)))

    with httpx.Client(timeout=httpx.Timeout(timeout_s, connect=5.0)) as client:
        response = client.post(
            f"{GOTENBERG_URL}/forms/libreoffice/convert",
            files=form_fields,
        )

    duration_ms = int((time.perf_counter() - start) * 1000)
    logger.info(
        f"[gotenberg] converting {input_format}→{output_format} "
        f"via /forms/libreoffice/convert in {duration_ms}ms"
    )

    if response.status_code != 200:
        raise RuntimeError(
            f"gotenberg: HTTP {response.status_code} for "
            f"{input_format}→{output_format}: {response.text[:500]}"
        )

    # Handle output
    if input_format == "pdf" and output_format == "docx":
        # Gotenberg returns PDF for PDF input; convert via pdftotext + docx.
        os.makedirs(GOTENBERG_SHARED_DIR, exist_ok=True)
        tmp_pdf = os.path.join(
            GOTENBERG_SHARED_DIR,
            f"got-tmp-{os.getpid()}-{time.time_ns()}.pdf",
        )
        try:
            with open(tmp_pdf, "wb") as f:
                for chunk in response.iter_bytes(chunk_size=8192):
                    f.write(chunk)
            _gotenberg_pdf_to_docx(tmp_pdf, output_path, timeout_s)
        finally:
            if os.path.exists(tmp_pdf):
                os.unlink(tmp_pdf)
    elif output_format in _GOTENBERG_DIRECT_FORMATS:
        with open(output_path, "wb") as f:
            for chunk in response.iter_bytes(chunk_size=8192):
                f.write(chunk)
    elif output_format == "pdf":
        _write_output(response, output_path)
    else:
        # Convert via intermediate PDF (txt, csv, html)
        os.makedirs(GOTENBERG_SHARED_DIR, exist_ok=True)
        tmp_pdf = os.path.join(
            GOTENBERG_SHARED_DIR,
            f"got-tmp-{os.getpid()}-{time.time_ns()}.pdf",
        )
        try:
            with open(tmp_pdf, "wb") as f:
                for chunk in response.iter_bytes(chunk_size=8192):
                    f.write(chunk)

            if output_format == "txt":
                _gotenberg_pdf_to_txt(tmp_pdf, output_path, timeout_s)
            elif output_format == "csv":
                _gotenberg_pdf_to_csv(tmp_pdf, output_path, timeout_s)
            elif output_format == "html":
                _gotenberg_pdf_to_html(tmp_pdf, output_path, timeout_s)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
        finally:
            if os.path.exists(tmp_pdf):
                os.unlink(tmp_pdf)

    logger.info(
        f"[gotenberg] done in {duration_ms}ms "
        f"({os.path.getsize(output_path)} bytes)"
    )


def _write_output(response: httpx.Response, output_path: str) -> None:
    try:
        with open(output_path, "wb") as f:
            for chunk in response.iter_bytes(chunk_size=8192):
                f.write(chunk)
    except Exception:
        if os.path.exists(output_path):
            os.unlink(output_path)
        raise


# Map from output format string to PIL format name
_PIL_FMT_MAP = {
    "png":  "PNG",
    "jpg":  "JPEG",
    "jpeg": "JPEG",
    "webp": "WEBP",
    "avif": "AVIF",
    "bmp":  "BMP",
    "tiff": "TIFF",
    "gif":  "GIF",
}


def _gotenberg_pdf_to_image(
    pdf_path: str,
    output_path: str,
    output_format: str,
    timeout_s: float,
) -> None:
    """Use pdf2image (pdftoppm + PIL) to render PDF pages as raster images."""
    from pdf2image import convert_from_path

    images = convert_from_path(
        pdf_path,
        dpi=150,
        first_page=1,
        last_page=1,
        fmt=output_format.upper(),
    )
    if not images:
        raise RuntimeError(f"[gotenberg] pdf2image produced no pages for {pdf_path}")
    img = images[0]
    img.save(output_path, format=_PIL_FMT_MAP[output_format])
    logger.info(f"[gotenberg] rendered PDF -> {output_format} via pdf2image")


def _gotenberg_pdf_to_txt(
    pdf_path: str,
    output_path: str,
    timeout_s: float,
) -> None:
    result = subprocess.run(
        ["pdftotext", "-layout", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
        check=True,
    )
    Path(output_path).write_bytes(result.stdout)


def _gotenberg_pdf_to_csv(
    pdf_path: str,
    output_path: str,
    timeout_s: float,
) -> None:
    result = subprocess.run(
        ["pdftotext", "-layout", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
        check=True,
    )
    text = result.stdout
    if isinstance(text, bytes):
        text = text.decode("utf-8", errors="replace")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    csv_rows = [",".join(f'"{f}"' for f in l.split()) for l in lines]
    Path(output_path).write_text("\n".join(csv_rows), encoding="utf-8")


def _gotenberg_pdf_to_html(
    pdf_path: str,
    output_path: str,
    timeout_s: float,
) -> None:
    result = subprocess.run(
        ["pdftotext", "-htmlmeta", "-q", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
    )
    if result.returncode == 0 and result.stdout.strip():
        Path(output_path).write_bytes(result.stdout)
    else:
        text_result = subprocess.run(
            ["pdftotext", "-layout", pdf_path, "-"],
            capture_output=True,
            timeout=timeout_s,
        )
        text = text_result.stdout
        if isinstance(text, bytes):
            text = text.decode("utf-8", errors="replace")
        lines = [
            "<p>{}</p>".format(p.strip())
            for p in text.split("\n\n")
            if p.strip()
        ]
        html = (
            "<!DOCTYPE html>\n<html><head>"
            "<meta charset=\"utf-8\"/>\n<body>\n"
            + "\n".join(lines)
            + "\n</body></html>"
        )
        Path(output_path).write_text(html, encoding="utf-8")


def _gotenberg_pdf_to_docx(
    pdf_path: str,
    output_path: str,
    timeout_s: float,
) -> None:
    """Convert PDF to DOCX via pdftotext + python-docx with sanitisation."""
    result = subprocess.run(
        ["pdftotext", "-layout", pdf_path, "-"],
        capture_output=True,
        timeout=timeout_s,
        check=True,
    )
    text = result.stdout

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for para_text in text.split(b"\n\n"):
        if isinstance(para_text, bytes):
            para_text = para_text.decode("utf-8", errors="replace")
        para_text = para_text.strip()
        if not para_text:
            continue
        # Remove control characters invalid in XML
        para_text = "".join(c for c in para_text if ord(c) >= 32 or c in "\t\n")
        if not para_text:
            continue
        para = doc.add_paragraph()
        para.add_run(para_text).font.size = Pt(11)

    doc.save(output_path)
